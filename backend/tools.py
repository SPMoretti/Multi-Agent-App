# tools.py

import os
import json
import base64
import psycopg2
import requests
import numpy as np
import pandas as pd
from datetime import datetime
from pydantic import BaseModel, Field
from crewai.tools import BaseTool
import fitz
from docx import Document

import storage as st_layer


# ─────────────────────────────────────────────────────────────
# 🔒 HELPER PATH
# Siempre trabaja con el path local (descargando desde Supabase si hace falta)
# ─────────────────────────────────────────────────────────────
def _safe_path(filename: str, username: str) -> str:
    """Devuelve el path local del archivo, descargándolo de Supabase si es necesario."""
    return st_layer.get_local_path(os.path.basename(filename), username)


def _read_text_from_file(path: str, ext: str) -> str:
    if ext == ".pdf":
        doc = fitz.open(path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    elif ext == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()


# ─────────────────────────────────────────────────────────────
# 🌐 SERP API
# ─────────────────────────────────────────────────────────────
class SerpApiTool(BaseTool):
    name: str = "Search the internet"
    description: str = "Busca informacion actualizada en internet usando SerpAPI."

    def _run(self, query: str) -> str:
        try:
            import serpapi
            today = datetime.now().strftime("%Y-%m-%d")
            current_year = datetime.now().year

            is_recent = any(str(y) in query for y in range(2024, current_year + 1)) or \
                        any(w in query.lower() for w in ["actual", "hoy", "último", "ultima", "reciente", "ahora", "hoy"])

            tbs_recent = f"cdr:1,cd_min:1/1/2024,cd_max:12/31/{current_year},sbd:1"
            tbs_value = tbs_recent if is_recent else "sbd:1"

            client = serpapi.Client(api_key=os.environ.get("SERPAPI_API_KEY", ""))
            results = client.search({
                "engine": "google",
                "q": query,
                "hl": "es",
                "gl": "ar",
                "num": 10,
                "tbs": tbs_value,
            })

            output = []
            for r in results.get("organic_results", []):
                output.append(r.get("title", ""))
                output.append(r.get("snippet", ""))
                output.append(r.get("link", ""))
                output.append("---")
            if not output:
                return "No se encontraron resultados."
            return f"[Busqueda: {today}]\n" + "\n".join(output)
        except Exception as e:
            return f"Error: {e}"


# ─────────────────────────────────────────────────────────────
# 📁 FILE TOOLS (con username)
# ─────────────────────────────────────────────────────────────
class FileWriterInput(BaseModel):
    filename: str = Field(description="Nombre del archivo, ej: datos.txt")
    content: str = Field(description="Contenido a escribir en el archivo")

class SafeFileWriterTool(BaseTool):
    name: str = "Write file safely"
    description: str = "Escribe archivos SOLO dentro de la carpeta del usuario. Soporta .txt, .csv, .json, .pdf, .docx"
    args_schema: type[BaseModel] = FileWriterInput
    username: str = ""

    def _run(self, filename: str, content: str) -> str:
        try:
            filename = os.path.basename(filename)
            ext = os.path.splitext(filename)[1].lower()

            if ext == ".docx":
                import io
                doc = Document()
                for line in content.split("\n"):
                    doc.add_paragraph(line)
                buf = io.BytesIO()
                doc.save(buf)
                file_bytes = buf.getvalue()
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            elif ext == ".pdf":
                import io
                doc = fitz.open()
                page = doc.new_page()
                page.insert_text((50, 50), content, fontsize=11)
                buf = io.BytesIO()
                doc.save(buf)
                doc.close()
                file_bytes = buf.getvalue()
                content_type = "application/pdf"

            else:
                file_bytes = content.encode("utf-8")
                content_type = "text/plain"

            st_layer.save_local_then_upload(filename, file_bytes, self.username, content_type)
            return f"Archivo '{filename}' guardado correctamente."

        except Exception as e:
            return f"Error escribiendo archivo: {e}"


# DESPUÉS
class FileReaderInput(BaseModel):
    filename: str = Field(description="Nombre del archivo a leer, ej: resultado.txt")
    query: str = Field(
        default="",
        description=(
            "Pregunta o contexto de lo que el usuario quiere saber. "
            "Completar siempre con la consulta original del usuario. "
            "Se usa para orientar la lectura si el archivo es grande."
        )
    )

class SafeFileReaderTool(BaseTool):
    name: str = "Read file safely"
    description: str = (
        "Lee el contenido de un archivo dentro de la carpeta del usuario. "
        "Soporta .txt, .csv, .json, .pdf, .docx. "
        "Completar siempre el campo 'query' con lo que el usuario quiere saber."
    )
    args_schema: type[BaseModel] = FileReaderInput
    username: str = ""

    def _run(self, filename: str, query: str = "") -> str:
        try:
            from config import CHAR_THRESHOLD
            filename = os.path.basename(filename)
            path = _safe_path(filename, self.username)

            if not os.path.exists(path):
                return f"❌ El archivo '{filename}' no existe."

            ext = os.path.splitext(filename)[1].lower()
            text = _read_text_from_file(path, ext)

            if not text.strip():
                return "El archivo está vacío."

            # Archivo chico → devolver completo
            if len(text) <= CHAR_THRESHOLD:
                return text

            # Archivo grande → RAG con query del usuario
            print(f"[FileReader] '{filename}' supera CHAR_THRESHOLD "
                  f"({len(text):,} chars) → derivando a RAG")

            rag_query_text = query.strip() if query.strip() else (
                "Resumí el contenido completo del documento incluyendo "
                "todos los puntos principales, datos, cifras y conclusiones."
            )

            from rag import rag_query
            result = rag_query(
                text=text,
                query=rag_query_text,
                filename=filename,
                filepath=path,
                username=self.username,
            )
            return (
                f"[Archivo grande — contenido recuperado por RAG]\n\n{result}\n\n"
                f"ℹ️ El archivo tiene {len(text):,} caracteres. "
                f"Para consultas más específicas usá 'Search file with RAG'."
            )

        except Exception as e:
            return f"Error leyendo archivo: {e}"


class ListFilesInput(BaseModel):
    dummy: str = Field(default="", description="No se necesitan argumentos")

class SafeListFilesTool(BaseTool):
    name: str = "List files in agent_files"
    description: str = "Lista todos los archivos disponibles dentro de la carpeta del usuario"
    args_schema: type[BaseModel] = ListFilesInput
    username: str = ""

    def _run(self, dummy: str = "") -> str:
        try:
            files = st_layer.list_files(self.username)
            if not files:
                return "La carpeta de archivos está vacía."
            return f"Archivos ({len(files)}):\n" + "\n".join(f"- {f}" for f in files)
        except Exception as e:
            return f"Error listando archivos: {e}"


class FileDeleterInput(BaseModel):
    filename: str = Field(description="Nombre del archivo a eliminar, ej: datos.txt")

class SafeFileDeleterTool(BaseTool):
    name: str = "Delete file safely"
    description: str = "Elimina un archivo dentro de la carpeta del usuario"
    args_schema: type[BaseModel] = FileDeleterInput
    username: str = ""

    def _run(self, filename: str) -> str:
        try:
            filename = os.path.basename(filename)
            deleted = st_layer.delete_file(filename, self.username)
            if deleted:
                return f"Archivo '{filename}' eliminado correctamente."
            return f"❌ El archivo '{filename}' no existe."
        except Exception as e:
            return f"Error eliminando archivo: {e}"


# ─────────────────────────────────────────────────────────────
# 🔍 RAG SEARCH TOOL (con username)
# ─────────────────────────────────────────────────────────────
class RAGSearchInput(BaseModel):
    filename: str = Field(description="Nombre del archivo a consultar, ej: informe.pdf")
    query: str = Field(description="Pregunta o consulta sobre el contenido del archivo")

class RAGSearchTool(BaseTool):
    name: str = "Search file with RAG"
    description: str = (
        "Busca información relevante dentro de un archivo usando RAG. "
        "Ideal para archivos grandes (.txt, .pdf, .docx .json)."
    )
    args_schema: type[BaseModel] = RAGSearchInput
    username: str = ""

    def _run(self, filename: str, query: str) -> str:
        try:
            from config import CHAR_THRESHOLD
            from rag import rag_query
            filename = os.path.basename(filename)
            path = _safe_path(filename, self.username)
            ext = os.path.splitext(filename)[1].lower()

            if not os.path.exists(path):
                return f"❌ El archivo '{filename}' no existe."

            if ext not in [".txt", ".pdf", ".docx", ".csv", ".json"]:
                return f"❌ Formato no soportado para RAG: {ext}"

            text = _read_text_from_file(path, ext)
            if not text.strip():
                return "El archivo está vacío o no contiene texto extraíble."

            if len(text) <= CHAR_THRESHOLD:
                from openai import OpenAI
                client = OpenAI()
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Respondé la pregunta basándote EXCLUSIVAMENTE en el contexto provisto."},
                        {"role": "user", "content": f"Contexto:\n\n{text}\n\nPregunta: {query}"}
                    ],
                    max_tokens=1024,
                    temperature=0.0,
                )
                return response.choices[0].message.content

            return rag_query(text, query, filename, path, self.username)

        except Exception as e:
            return f"Error en RAG: {e}"


# ─────────────────────────────────────────────────────────────
# 🖼️ IMAGE TOOLS (con username)
# ─────────────────────────────────────────────────────────────
class ImageDescriptorInput(BaseModel):
    filename: str = Field(description="Nombre de la imagen, ej: foto.jpg")

class ImageDescriptorTool(BaseTool):
    name: str = "Describe image"
    description: str = "Describe el contenido de una imagen usando visión IA. Soporta .jpg, .jpeg, .png, .gif"
    args_schema: type[BaseModel] = ImageDescriptorInput
    username: str = ""

    def _run(self, filename: str) -> str:
        try:
            from openai import OpenAI
            filename = os.path.basename(filename)
            path = _safe_path(filename, self.username)
            ext = os.path.splitext(filename)[1].lower()

            if ext not in [".jpg", ".jpeg", ".png", ".gif"]:
                return f"❌ Formato no soportado: {ext}"

            if not os.path.exists(path):
                return f"❌ El archivo '{filename}' no existe."

            with open(path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode("utf-8")

            mime = "image/jpeg" if ext in [".jpg", ".jpeg"] else f"image/{ext[1:]}"
            client = OpenAI()
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{image_data}"}},
                            {"type": "text", "text": "Describí detalladamente el contenido de esta imagen en español."}
                        ]
                    }
                ],
                max_tokens=512,
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error describiendo imagen: {e}"


class ImageGeneratorInput(BaseModel):
    prompt: str = Field(description="Descripción en inglés de la imagen a generar")
    filename: str = Field(description="Nombre del archivo .png a guardar")

class ImageGeneratorTool(BaseTool):
    name: str = "Generate image"
    description: str = "Genera una imagen con DALL-E 2 y la guarda en la carpeta del usuario"
    args_schema: type[BaseModel] = ImageGeneratorInput
    username: str = ""

    def _run(self, prompt: str, filename: str) -> str:
        try:
            from openai import OpenAI
            client = OpenAI()

            response = client.images.generate(
                model="dall-e-2",
                prompt=prompt,
                size="256x256",
                n=1,
                response_format="b64_json",
            )

            image_b64 = response.data[0].b64_json
            image_bytes = base64.b64decode(image_b64)

            filename = os.path.basename(filename)
            if not filename.lower().endswith(".png"):
                filename = os.path.splitext(filename)[0] + ".png"

            st_layer.save_local_then_upload(filename, image_bytes, self.username, "image/png")
            return f"✅ Imagen generada y guardada como '{filename}'."

        except Exception as e:
            return f"Error generando imagen: {e}"


# ─────────────────────────────────────────────────────────────
# 📊 DATA ANALYSIS TOOL
# ─────────────────────────────────────────────────────────────
class DataAnalysisInput(BaseModel):
    data: str = Field(description="Datos en formato JSON (lista de objetos) o CSV como string")
    operation: str = Field(description="Operacion a realizar: stats | outliers | correlation | regression")
    columns: list[str] | None = Field(default=None, description="Columnas a analizar (opcional)")

class DataAnalysisTool(BaseTool):
    name: str = "Analyze data"
    description: str = "Realiza análisis estadístico sobre datos estructurados."
    args_schema: type[BaseModel] = DataAnalysisInput

    def _load_data(self, data: str) -> pd.DataFrame:
        try:
            return pd.DataFrame(json.loads(data))
        except Exception:
            from io import StringIO
            return pd.read_csv(StringIO(data))

    def _run(self, data: str, operation: str, columns=None) -> str:
        try:
            df = self._load_data(data)
            if df.empty:
                return "❌ Dataset vacío."
            if columns:
                df = df[columns]
            numeric_df = df.select_dtypes(include=[np.number])
            if numeric_df.empty:
                return "❌ No hay columnas numéricas para analizar."

            if operation == "stats":
                return json.dumps(numeric_df.describe().to_dict(), indent=2)
            elif operation == "outliers":
                outliers = {}
                for col in numeric_df.columns:
                    q1, q3 = numeric_df[col].quantile([0.25, 0.75])
                    iqr = q3 - q1
                    mask = (numeric_df[col] < q1 - 1.5 * iqr) | (numeric_df[col] > q3 + 1.5 * iqr)
                    outliers[col] = numeric_df[col][mask].tolist()
                return json.dumps(outliers, indent=2)
            elif operation == "correlation":
                return json.dumps(numeric_df.corr().to_dict(), indent=2)
            elif operation == "regression":
                if not columns or len(columns) != 2:
                    return "❌ Para regresión se necesitan exactamente 2 columnas."
                x, y = numeric_df[columns[0]], numeric_df[columns[1]]
                slope, intercept = np.polyfit(x, y, 1)
                return json.dumps({"slope": slope, "intercept": intercept}, indent=2)
            else:
                return f"❌ Operación no soportada: {operation}"
        except Exception as e:
            return f"Error en análisis de datos: {e}"


# ─────────────────────────────────────────────────────────────
# 🗄️ POSTGRES (con schema por usuario)
# ─────────────────────────────────────────────────────────────
def _pg_conn(username: str = ""):
    host     = os.getenv("PG_HOST")
    port     = os.getenv("PG_PORT", "5432")
    database = os.getenv("PG_DB")
    user     = os.getenv("PG_USER")
    password = os.getenv("PG_PASSWORD")

    print(f"[DB] Conectando a {host}:{port}/{database} como {user}")

    if not all([host, database, user, password]):
        missing = [k for k, v in {"PG_HOST": host, "PG_DB": database, "PG_USER": user, "PG_PASSWORD": password}.items() if not v]
        raise ValueError(f"Faltan variables de entorno: {missing}")

    conn = psycopg2.connect(
        host=host,
        port=int(port),
        database=database,
        user=user,
        password=password,
        sslmode="require",
        options="-c client_encoding=UTF8",
    )

    # Establecer search_path al schema del usuario si se provee
    if username:
        from config import get_schema_name
        schema = get_schema_name(username)
        cur = conn.cursor()
        cur.execute(f'SET search_path TO "{schema}", public')
        cur.close()
        conn.commit()

    return conn


def get_schema(username: str = "") -> str:
    """
    Obtiene el schema de la base de datos.
    Si se provee username, usa su schema personal; sino usa 'public'.
    """
    try:
        conn = _pg_conn(username)
        conn.set_client_encoding("UTF8")
        cur = conn.cursor()

        if username:
            from config import get_schema_name
            schema_name = get_schema_name(username)
        else:
            schema_name = "public"

        cur.execute("""
            SELECT table_name, column_name, data_type
            FROM information_schema.columns
            WHERE table_schema = %s
            ORDER BY table_name;
        """, (schema_name,))
        rows = cur.fetchall()
        cur.close()
        conn.close()

        if not rows:
            print(f"[DB] Schema vacío: no se encontraron tablas en schema '{schema_name}'")
            return f"No se encontraron tablas en el schema '{schema_name}'."

        schema = {}
        for table, col, dtype in rows:
            schema.setdefault(table, []).append(f"{col} ({dtype})")

        lines = []
        for table, cols in schema.items():
            lines.append(f"Tabla {table}:")
            for c in cols:
                lines.append(f"  - {c}")

        result = "\n".join(lines)
        print(f"[DB] Schema obtenido:\n{result}")
        return result

    except Exception as e:
        import traceback
        error = f"Error obteniendo schema: {e}\n{traceback.format_exc()}"
        print(f"[DB] {error}")
        return error


class SafePostgresTool(BaseTool):
    name: str = Field(default="Query PostgreSQL safely")
    description: str = Field(default="Ejecuta consultas SQL SELECT seguras sobre PostgreSQL.")
    username: str = ""

    def _run(self, query: str) -> str:
        try:
            q = query.strip().lower()
            if not q.startswith("select"):
                return "❌ Solo se permiten queries SELECT."
            if any(w in q for w in ["drop", "delete", "update", "insert"]):
                return "❌ Query peligrosa bloqueada."

            conn = _pg_conn(self.username)
            conn.set_client_encoding("UTF8")
            cur = conn.cursor()
            cur.execute(query)
            rows = cur.fetchmany(50)
            cols = [desc[0] for desc in cur.description]
            result = [dict(zip(cols, row)) for row in rows]
            cur.close()
            conn.close()
            return json.dumps(result)
        except Exception as e:
            import traceback
            return f"SQL_ERROR: {str(e)}\n{traceback.format_exc()}"


# ─────────────────────────────────────────────────────────────
# FACTORIES (reciben username)
# ─────────────────────────────────────────────────────────────
def get_file_tools(username: str = "") -> list:
    from config import get_files_dir
    if username:
        os.makedirs(get_files_dir(username), exist_ok=True)
    return [
        SafeFileReaderTool(username=username),
        SafeFileWriterTool(username=username),
        SafeListFilesTool(username=username),
        SafeFileDeleterTool(username=username),
        ImageDescriptorTool(username=username),
        RAGSearchTool(username=username),
        DataAnalysisTool(),
    ]

def get_image_tools(username: str = "") -> list:
    return [ImageGeneratorTool(username=username)]